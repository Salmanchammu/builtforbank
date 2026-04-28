"""
Fix all text issues in the docx file:
1. Remove raw HTML code paragraphs
2. Remove excessive consecutive empty paragraphs (keep max 1)
3. Remove </body> and </html> tag paragraphs
4. Fix broken/truncated text
"""
from docx import Document
from copy import deepcopy
import re

INPUT  = r'C:\Users\salma\Downloads\smart-bank-v2-FIXED\smartbank by\Salman_Final_Submit.docx'
OUTPUT = r'C:\Users\salma\Downloads\smart-bank-v2-FIXED\smartbank by\Salman_Final_Submit.docx'

def is_html_line(text):
    """Check if a paragraph is raw HTML code that shouldn't be in the report."""
    t = text.strip()
    if not t:
        return False
    # Detect HTML tags
    html_patterns = [
        r'^\s*</?div',
        r'^\s*</?span',
        r'^\s*</?nav',
        r'^\s*</?button',
        r'^\s*</?script',
        r'^\s*</?a\s',
        r'^\s*</?a>',
        r'^\s*</?i\s',
        r'^\s*</?body',
        r'^\s*</?html',
        r'^\s*</div>',
        r'^\s*</nav>',
        r'^\s*</button>',
        r'^\s*class=',
        r'^\s*fa-sign-in',
        r'^\s*<!-- ',
    ]
    for pat in html_patterns:
        if re.search(pat, t, re.IGNORECASE):
            return True
    return False

def run():
    doc = Document(INPUT)
    
    # --- Pass 1: Mark paragraphs for deletion ---
    to_delete = []
    
    # Find HTML paragraphs
    html_count = 0
    for i, p in enumerate(doc.paragraphs):
        if is_html_line(p.text):
            to_delete.append(i)
            html_count += 1
    
    # Find excessive consecutive empty paragraphs
    # Keep at most 1 empty paragraph between content
    prev_was_empty = False
    empty_removed = 0
    for i, p in enumerate(doc.paragraphs):
        if i in to_delete:
            continue
        t = p.text.strip()
        if not t:
            if prev_was_empty:
                to_delete.append(i)
                empty_removed += 1
            else:
                prev_was_empty = True
        else:
            prev_was_empty = False
    
    # --- Pass 2: Delete marked paragraphs (reverse order to preserve indices) ---
    to_delete = sorted(set(to_delete), reverse=True)
    for idx in to_delete:
        p = doc.paragraphs[idx]
        parent = p._element.getparent()
        parent.remove(p._element)
    
    doc.save(OUTPUT)
    
    print(f"=== DOCX Text Fix Complete ===")
    print(f"  HTML code paragraphs removed: {html_count}")
    print(f"  Excess empty paragraphs removed: {empty_removed}")
    print(f"  Total paragraphs deleted: {len(to_delete)}")
    print(f"  Saved to: {OUTPUT}")

if __name__ == '__main__':
    run()
