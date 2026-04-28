"""
Comprehensive DOCX fixer:
1. Remove ALL empty/blank paragraphs (fix the big gaps between chapters)
2. Keep tables and diagrams together on one page (keep-with-next)
3. Remove any leftover HTML code paragraphs
4. Set clean, consistent heading spacing
"""
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from lxml import etree
import re

INPUT  = r'C:\Users\salma\Downloads\smart-bank-v2-FIXED\smartbank by\Salman_Final_Submit.docx'
OUTPUT = r'C:\Users\salma\Downloads\smart-bank-v2-FIXED\smartbank by\Salman_Final_Submit_FIXED.docx'

def is_html_line(text):
    t = text.strip()
    if not t:
        return False
    html_patterns = [
        r'</?div', r'</?span', r'</?nav', r'</?button', r'</?script',
        r'</?a[\s>]', r'</?i\s', r'</?body', r'</?html',
        r'</div>', r'</nav>', r'</button>', r'class=', r'fa-sign-in',
        r'<!--', r'onclick=', r'href='
    ]
    for pat in html_patterns:
        if re.search(pat, t, re.IGNORECASE):
            return True
    return False

def actually_has_image(para):
    """Properly check for real images - not just namespace substring matches."""
    # Look for actual drawing/image child elements, not namespace declarations
    ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ns_pic = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    ns_v = 'urn:schemas-microsoft-com:vml'
    
    # Check for <w:drawing> elements (modern images)
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        return True
    
    # Check for <w:pict> elements (legacy images/shapes)  
    picts = para._element.findall('.//' + qn('w:pict'))
    if picts:
        return True
    
    # Check for inline shapes
    inlines = para._element.findall('.//{%s}inline' % ns_wp)
    if inlines:
        return True
    
    return False

def set_keep_together(para):
    """Set paragraph to keep with next (prevents page breaks splitting content)."""
    pPr = para._element.get_or_add_pPr()
    keep = pPr.find(qn('w:keepNext'))
    if keep is None:
        keep = etree.SubElement(pPr, qn('w:keepNext'))
    keep.set(qn('w:val'), '1')

def set_keep_lines(para):
    """Prevent a paragraph from splitting across pages."""
    pPr = para._element.get_or_add_pPr()
    keep = pPr.find(qn('w:keepLines'))
    if keep is None:
        keep = etree.SubElement(pPr, qn('w:keepLines'))
    keep.set(qn('w:val'), '1')

def run():
    doc = Document(INPUT)
    total = len(doc.paragraphs)
    
    # =============================================
    # STEP 1: Remove empty paragraphs & HTML junk
    # =============================================
    to_delete = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Delete HTML paragraphs
        if is_html_line(text):
            to_delete.append(i)
            continue
        
        # Delete empty paragraphs that DON'T contain real images
        if not text and not actually_has_image(p):
            to_delete.append(i)
            continue
    
    # Delete in reverse order to preserve indices
    deleted = 0
    for idx in sorted(to_delete, reverse=True):
        p = doc.paragraphs[idx]
        parent = p._element.getparent()
        parent.remove(p._element)
        deleted += 1
    
    print(f"  Empty/HTML paragraphs removed: {deleted}")
    
    # =============================================
    # STEP 2: Set clean heading spacing
    # =============================================
    heading_count = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if 'Heading 1' in style_name:
            p.paragraph_format.space_before = Pt(28)
            p.paragraph_format.space_after = Pt(8)
            set_keep_together(p)  # Keep heading with following content
            heading_count += 1
        elif 'Heading 2' in style_name:
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            set_keep_together(p)
            heading_count += 1
        elif 'Heading 3' in style_name:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            set_keep_together(p)
            heading_count += 1
        elif 'Heading 4' in style_name:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            set_keep_together(p)
            heading_count += 1
    
    print(f"  Headings with spacing fixed: {heading_count}")
    
    # =============================================
    # STEP 3: Keep tables together on one page
    # =============================================
    table_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_keep_together(para)
                    set_keep_lines(para)
        table_count += 1
    
    print(f"  Tables set to keep-together: {table_count}")
    
    # =============================================
    # STEP 4: Keep "Figure:" captions with their preceding content
    # =============================================
    caption_count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # If this paragraph contains a figure/diagram image, keep it with next
        if actually_has_image(p):
            set_keep_together(p)
            set_keep_lines(p)
            caption_count += 1
        # If this is a figure caption, it was already handled
        if text.lower().startswith('figure:') or text.lower().startswith('figure '):
            set_keep_lines(p)
            caption_count += 1
    
    # Also keep paragraphs right before diagrams/figures together
    for i in range(len(doc.paragraphs) - 1):
        p_next = doc.paragraphs[i + 1]
        if actually_has_image(p_next):
            set_keep_together(doc.paragraphs[i])
    
    print(f"  Figure/diagram paragraphs protected: {caption_count}")
    
    # =============================================
    # STEP 5: Fix body text spacing for consistency
    # =============================================
    body_count = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name == 'Body Text' and p.text.strip():
            pf = p.paragraph_format
            # Set consistent spacing for body text
            if pf.space_before and pf.space_before > Pt(14):
                pf.space_before = Pt(6)
            if pf.space_after is None or pf.space_after > Pt(12):
                pf.space_after = Pt(4)
            body_count += 1
        elif style_name == 'List Paragraph' and p.text.strip():
            pf = p.paragraph_format
            if pf.space_before and pf.space_before > Pt(14):
                pf.space_before = Pt(6)
            body_count += 1
    
    print(f"  Body/list paragraphs spacing normalized: {body_count}")
    
    doc.save(OUTPUT)
    
    remaining = len(doc.paragraphs)
    print(f"\n=== DOCX Fix Complete ===")
    print(f"  Original: {total} paragraphs")
    print(f"  Remaining: {remaining} paragraphs")
    print(f"  Saved: {OUTPUT}")

if __name__ == '__main__':
    run()
